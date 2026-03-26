"""
FileManager — lecture, renommage et conversion de documents Word pour CV et lettres de motivation.

Fonctionnalités :
  - Lecture du contenu texte d'un fichier .docx
  - Construction d'un chemin de sortie unique basé sur le nom du poste
  - Sauvegarde du contenu modifié dans un nouveau .docx (sans écraser l'original)
  - Conversion .docx → PDF via docx2pdf (primary) ou win32com (fallback)
  - Pipeline complet : save_docx + convert_to_pdf en une seule opération

Usage :
    from linkedin.utils.file_manager import FileManager

    # Lecture
    content = FileManager.read_docx("C:/docs/lettre_motivation.docx")

    # Pipeline complet (appelé après AIManager)
    paths = FileManager.process_cover_letter(
        original_docx_path="C:/docs/lettre_motivation.docx",
        new_content=improved_text,
        position_name="Data Scientist",
        output_dir="C:/docs/candidatures/",
        convert_pdf=True,
    )
    print(paths["docx"])  # C:/docs/candidatures/lettre_motivation_DataScientist.docx
    print(paths["pdf"])   # C:/docs/candidatures/lettre_motivation_DataScientist.pdf
"""
import logging
import os
import re

logger = logging.getLogger(__name__)


class FileManager:
    """
    Utilitaires statiques pour gérer les fichiers de candidature (CV, lettres de motivation).

    Toutes les méthodes sont statiques ou de classe : aucune instance n'est nécessaire.
    Le fichier original n'est jamais modifié — toute écriture produit un nouveau fichier.
    """

    # ──────────────────────────────────────────────────────────────
    # Lecture
    # ──────────────────────────────────────────────────────────────

    @staticmethod
    def read_docx(filepath: str) -> str:
        """
        Extrait le contenu texte d'un fichier .docx.

        Chaque paragraphe est séparé par un saut de ligne.
        Les paragraphes vides sont préservés pour conserver la mise en page.

        Args:
            filepath: Chemin absolu vers le fichier .docx.

        Returns:
            Contenu textuel complet du document.

        Raises:
            FileNotFoundError: Si le fichier n'existe pas.
            ValueError: Si l'extension n'est pas .docx.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx est requis pour lire les fichiers .docx. "
                "Installez-le avec : pip install python-docx"
            ) from exc

        filepath = os.path.abspath(filepath)
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Fichier introuvable : {filepath}")
        if not filepath.lower().endswith(".docx"):
            raise ValueError(f"Extension non supportée (attendu .docx) : {filepath}")

        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)

    # ──────────────────────────────────────────────────────────────
    # Gestion des chemins de sortie
    # ──────────────────────────────────────────────────────────────

    @staticmethod
    def make_output_path(
        original_path: str,
        position_name: str,
        extension: str = ".docx",
        output_dir: str | None = None,
    ) -> str:
        """
        Construit un chemin de sortie unique à partir du nom de l'original et du poste.

        Exemple :
            original_path = "C:/docs/lettre_motivation.docx"
            position_name = "Data Scientist"
            → "C:/docs/lettre_motivation_DataScientist.docx"

        Si le fichier existe déjà, un suffixe numérique est ajouté :
            → "C:/docs/lettre_motivation_DataScientist_1.docx"

        Args:
            original_path:  Chemin du fichier source (utilisé pour le nom de base et le dossier).
            position_name:  Nom du poste à intégrer dans le nom de fichier (sera nettoyé).
            extension:      Extension du fichier de sortie (".docx" ou ".pdf").
            output_dir:     Dossier de destination. Si None, utilise le dossier de l'original.

        Returns:
            Chemin absolu garanti unique (le fichier n'existe pas encore).
        """
        original_path = os.path.abspath(original_path)
        base_name = os.path.splitext(os.path.basename(original_path))[0]

        # Nettoie le nom du poste : caractères non alphanumériques → underscore
        slug = re.sub(r"[^\w]", "_", position_name).strip("_")
        # Supprime les underscores consécutifs
        slug = re.sub(r"_+", "_", slug)

        # Dossier de sortie
        dest_dir = output_dir if output_dir else os.path.dirname(original_path)
        os.makedirs(dest_dir, exist_ok=True)

        # Chemin candidat sans suffixe
        candidate = os.path.join(dest_dir, f"{base_name}_{slug}{extension}")

        # Ajoute _1, _2, … jusqu'à trouver un chemin libre
        counter = 1
        while os.path.exists(candidate):
            candidate = os.path.join(dest_dir, f"{base_name}_{slug}_{counter}{extension}")
            counter += 1

        return candidate

    # ──────────────────────────────────────────────────────────────
    # Écriture
    # ──────────────────────────────────────────────────────────────

    @staticmethod
    def save_docx(content: str, output_path: str) -> str:
        """
        Écrit un contenu texte dans un nouveau fichier .docx.

        Les sauts de ligne séparent les paragraphes.
        Les lignes vides produisent des paragraphes vides (espacement visuel préservé).

        Args:
            content:     Texte à écrire (paragraphes séparés par \\n).
            output_path: Chemin absolu du nouveau fichier (ne doit pas exister).

        Returns:
            Le chemin ``output_path`` écrit.

        Raises:
            FileExistsError: Si output_path existe déjà (utilisez make_output_path).
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx est requis. Installez-le avec : pip install python-docx"
            ) from exc

        output_path = os.path.abspath(output_path)
        if os.path.exists(output_path):
            raise FileExistsError(
                f"Le fichier existe déjà : {output_path}. "
                "Utilisez make_output_path() pour obtenir un chemin unique."
            )

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)

        doc.save(output_path)
        logger.debug("Fichier .docx sauvegardé : %s", output_path)
        return output_path

    # ──────────────────────────────────────────────────────────────
    # Conversion PDF
    # ──────────────────────────────────────────────────────────────

    @staticmethod
    def convert_to_pdf(docx_path: str, pdf_path: str | None = None) -> str:
        """
        Convertit un fichier .docx en PDF.

        Stratégie de conversion (essayées dans l'ordre) :
          1. ``docx2pdf.convert()`` — requiert Microsoft Word installé sur Windows.
          2. ``win32com.client`` — automatisation COM Word (fallback Windows).

        Args:
            docx_path: Chemin absolu vers le .docx source.
            pdf_path:  Chemin absolu pour le PDF de sortie.
                       Si None, remplace l'extension .docx par .pdf dans le même dossier.

        Returns:
            Chemin absolu vers le PDF généré.

        Raises:
            FileNotFoundError: Si docx_path n'existe pas.
            RuntimeError: Si les deux méthodes de conversion échouent.
        """
        docx_path = os.path.abspath(docx_path)
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Fichier source introuvable : {docx_path}")

        if pdf_path is None:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        pdf_path = os.path.abspath(pdf_path)
        os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)

        primary_exc: Exception | None = None

        # ── Méthode 1 : docx2pdf ─────────────────────────────────
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            logger.debug("PDF généré via docx2pdf : %s", pdf_path)
            return pdf_path
        except Exception as exc:
            primary_exc = exc
            logger.debug("docx2pdf a échoué : %s", exc)

        # ── Méthode 2 : win32com (COM Word) ─────────────────────
        try:
            import win32com.client  # type: ignore[import]
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
                doc.Close()
                logger.debug("PDF généré via win32com : %s", pdf_path)
                return pdf_path
            finally:
                word.Quit()
        except Exception as fallback_exc:
            raise RuntimeError(
                "La conversion PDF a échoué avec les deux méthodes.\n"
                f"  docx2pdf : {primary_exc}\n"
                f"  win32com : {fallback_exc}\n\n"
                "Assurez-vous que Microsoft Word est installé, "
                "ou installez docx2pdf : pip install docx2pdf"
            ) from fallback_exc

    # ──────────────────────────────────────────────────────────────
    # Pipeline complet
    # ──────────────────────────────────────────────────────────────

    @classmethod
    def process_cover_letter(
        cls,
        original_docx_path: str,
        new_content: str,
        position_name: str,
        output_dir: str | None = None,
        convert_pdf: bool = True,
    ) -> dict[str, str]:
        """
        Pipeline complet : sauvegarde la lettre améliorée en .docx et (optionnel) en PDF.

        Cette méthode n'écrase jamais le fichier original.
        Le contenu ``new_content`` doit déjà être généré (ex. par AIManager).

        Args:
            original_docx_path: Chemin du .docx original (utilisé comme modèle de nommage).
            new_content:        Texte amélioré de la lettre de motivation.
            position_name:      Nom du poste pour construire le nom du fichier de sortie.
            output_dir:         Dossier de destination. Si None, utilise le dossier de l'original.
            convert_pdf:        Si True, génère également un PDF à côté du .docx.

        Returns:
            Dict avec les chemins générés :
            {
                "docx": "<chemin absolu du nouveau .docx>",
                "pdf":  "<chemin absolu du PDF>"  # chaîne vide si convert_pdf=False
            }
        """
        # ── Étape 1 : construire le chemin .docx ─────────────────
        docx_output = cls.make_output_path(
            original_path=original_docx_path,
            position_name=position_name,
            extension=".docx",
            output_dir=output_dir,
        )

        # ── Étape 2 : sauvegarder le .docx ───────────────────────
        cls.save_docx(new_content, docx_output)
        print(f"  📄 Lettre sauvegardée : {docx_output}")

        # ── Étape 3 : convertir en PDF (optionnel) ────────────────
        pdf_output = ""
        if convert_pdf:
            pdf_candidate = cls.make_output_path(
                original_path=original_docx_path,
                position_name=position_name,
                extension=".pdf",
                output_dir=output_dir,
            )
            try:
                pdf_output = cls.convert_to_pdf(docx_output, pdf_candidate)
                print(f"  📋 PDF généré       : {pdf_output}")
            except RuntimeError as exc:
                logger.warning("Conversion PDF échouée : %s", exc)
                print(f"  ⚠️  Conversion PDF échouée (voir les logs) : {exc}")

        return {"docx": docx_output, "pdf": pdf_output}
